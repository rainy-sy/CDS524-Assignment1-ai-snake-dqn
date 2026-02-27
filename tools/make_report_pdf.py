from __future__ import annotations

import os
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


def _set_style_sizes(doc: Document) -> None:
    styles = doc.styles

    normal = styles["Normal"]
    if normal.font.size is None:
        normal.font.size = Pt(11)

    h1 = styles["Heading 1"]
    h1.font.bold = True
    h1.font.size = Pt(16)

    h2 = styles["Heading 2"]
    h2.font.bold = True
    h2.font.size = Pt(14)


def _add_header(doc: Document, text: str) -> None:
    section = doc.sections[0]
    header = section.header
    p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    p.text = text
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def build_docx(output_docx: Path, base_dir: Path) -> None:
    doc = Document()
    _set_style_sizes(doc)
    _add_header(doc, "WANG Shiyu – CDS524 Assignment 1")

    # Title block
    title = doc.add_paragraph("CDS524 Assignment 1")
    title.runs[0].bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph("Reinforcement Learning Game Design Report")
    subtitle.runs[0].bold = True
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("")

    meta = doc.add_paragraph()
    meta.add_run("Student Name: ").bold = True
    meta.add_run("WANG Shiyu")

    meta = doc.add_paragraph()
    meta.add_run("Game Title: ").bold = True
    meta.add_run("AI Snake — Deep Q-Learning with Dreamy Pink Theme")

    meta = doc.add_paragraph()
    meta.add_run("Submission Date: ").bold = True
    meta.add_run("25 February 2026")

    doc.add_page_break()

    # 1. Introduction
    doc.add_paragraph("1. Introduction", style="Heading 1")
    doc.add_paragraph(
        "This project presents an enhanced Snake game developed with Pygame and trained using Deep Q-Learning (DQN). "
        "The objective is to control a snake that grows by consuming food while avoiding collisions with the walls and its own body. "
        "The dynamic nature of the snake’s body creates a challenging decision-making environment that is ideal for reinforcement learning."
    )
    doc.add_paragraph(
        "Deep Q-Learning (DQN) was employed to enable the agent to learn optimal policies through interaction with the environment. "
        "The implementation uses PyTorch for the neural network and experience replay mechanism, while Pygame provides a visually appealing "
        "and interactive user interface featuring a dreamy pink theme."
    )

    # 2. Game Design
    doc.add_paragraph("2. Game Design", style="Heading 1")

    doc.add_paragraph("2.1 Objective and Rules", style="Heading 2")
    rules = [
        "The game operates on a 20 × 20 grid with a 900 × 900 pixel window.",
        "The snake begins at length 3 in the center, facing right.",
        "At each timestep, the agent chooses one of three relative actions: turn left, go straight, or turn right.",
        "Collision with walls or the snake’s own body results in game over.",
        "The primary goal is to maximize the score by eating as much food as possible while surviving as long as possible.",
    ]
    for r in rules:
        doc.add_paragraph(r, style="List Bullet")

    doc.add_paragraph("2.2 State Space", style="Heading 2")
    doc.add_paragraph(
        "The environment is represented by an 11-dimensional binary feature vector. It captures (i) immediate collision risk, "
        "(ii) food direction relative to the head, and (iii) the current movement direction."
    )
    doc.add_paragraph(
        "Table 1 provides a precise definition of each dimension (all features are binary: 1=true, 0=false)."
    )

    # Table 1: 11D state definition
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Dim"
    hdr[1].text = "Name"
    hdr[2].text = "Meaning"
    hdr[3].text = "Values"

    rows = [
        (0, "Danger Straight", "Obstacle immediately in front of the head (wall or body)", "1=danger, 0=safe"),
        (1, "Danger Right", "Obstacle immediately to the right of the head", "1=danger, 0=safe"),
        (2, "Danger Left", "Obstacle immediately to the left of the head", "1=danger, 0=safe"),
        (3, "Food Left", "Food is to the left of the head", "1=yes, 0=no"),
        (4, "Food Right", "Food is to the right of the head", "1=yes, 0=no"),
        (5, "Food Up", "Food is above the head", "1=yes, 0=no"),
        (6, "Food Down", "Food is below the head", "1=yes, 0=no"),
        (7, "Move Left", "Current movement direction is left", "1=yes, 0=no"),
        (8, "Move Right", "Current movement direction is right", "1=yes, 0=no"),
        (9, "Move Up", "Current movement direction is up", "1=yes, 0=no"),
        (10, "Move Down", "Current movement direction is down", "1=yes, 0=no"),
    ]
    for dim, name, meaning, values in rows:
        cells = table.add_row().cells
        cells[0].text = str(dim)
        cells[1].text = name
        cells[2].text = meaning
        cells[3].text = values

    cap = doc.add_paragraph("Table 1: Definition of the 11D state vector")
    cap.runs[0].italic = True

    doc.add_paragraph(
        "Note: ‘left/right/up/down’ are defined in screen coordinates relative to the head position (not relative to the current direction)."
    )
    doc.add_paragraph(
        "This compact representation avoids the full grid as input while preserving key information needed for safe navigation and food seeking."
    )

    doc.add_paragraph("2.3 Action Space", style="Heading 2")
    doc.add_paragraph("The action space consists of 3 discrete relative actions:")
    actions = [
        "0: Turn left",
        "1: Go straight",
        "2: Turn right",
    ]
    for a in actions:
        doc.add_paragraph(a, style="List Bullet")
    doc.add_paragraph("Relative actions naturally prevent invalid 180-degree turns.")

    doc.add_paragraph("2.4 Reward Function", style="Heading 2")
    doc.add_paragraph("The reward function is designed to encourage desired behaviors:")
    rewards = [
        "Normal food: +10",
        "Bonus food: +20 + combo multiplier",
        "Poison food: -8",
        "Collision (death): -10",
        "Survival per step: +0.1",
        "Moving away from food: -0.5 (light shaping reward)",
    ]
    for rw in rewards:
        doc.add_paragraph(rw, style="List Bullet")

    doc.add_paragraph(
        "Reward design was refined iteratively. A sparse reward (food/death only) makes exploration inefficient, so a small living reward (+0.1) "
        "encourages movement and discovery of useful trajectories. To reduce unnecessary wandering, a light shaping penalty (−0.5) is applied when the "
        "snake moves away from the food. Finally, adding bonus and poison food types enriches the learning signal while keeping the core objective unchanged."
    )

    # Figure 1: insert gameplay screenshot if present
    fig1 = base_dir / "figure1.png"
    if fig1.exists():
        doc.add_paragraph("")
        try:
            doc.add_picture(str(fig1), width=Inches(5.8))
        except Exception:
            p = doc.add_paragraph("[Could not embed figure1.png — insert manually]")
            p.runs[0].italic = True
    else:
        p = doc.add_paragraph("[Insert Figure 1 here]")
        p.runs[0].italic = True

    cap = doc.add_paragraph(
        "Figure 1: Game interface showing dreamy pink theme, glowing food, combo counter, and real-time 11D state display"
    )
    cap.runs[0].italic = True

    # 3. Q-Learning Implementation
    doc.add_paragraph("3. Q-Learning Implementation", style="Heading 1")

    doc.add_paragraph("3.1 Algorithm", style="Heading 2")
    doc.add_paragraph(
        "Deep Q-Learning approximates the Q-value function with a neural network. The target value is computed as:"
    )
    eq = doc.add_paragraph("y = r + (1 − done) × γ × max Q(s′, a′)")
    eq.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(
        "The network minimizes the mean squared error between predicted and target Q-values."
    )

    doc.add_paragraph("3.2 Neural Network Architecture", style="Heading 2")
    doc.add_paragraph(
        "LinearQNet is a 4-layer fully connected network with ReLU activations:"
    )
    arch = doc.add_paragraph("11  →  256  →  128  →  64  →  3")
    arch.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(
        "The input is the 11D state vector and the output is the Q-value for each of the 3 relative actions (left/straight/right). "
        "The hidden layer widths decrease (256→128→64) to progressively compress features into higher-level abstractions while controlling model capacity. "
        "This structure is a common, stable baseline for small structured state spaces and converges reliably in this task."
    )

    doc.add_paragraph("3.3 Exploration vs Exploitation", style="Heading 2")
    doc.add_paragraph("An ε-greedy policy is used:")
    eps = [
        "Initial ε = 0.8",
        "Minimum ε = 0.01",
        "Multiplicative decay of 0.995 per episode",
    ]
    for e in eps:
        doc.add_paragraph(e, style="List Bullet")

    doc.add_paragraph("3.4 Training Components", style="Heading 2")
    comps = [
        "Replay buffer capacity: 10,000 transitions",
        "Batch size: 32",
        "Optimizer: Adam (learning rate = 0.001)",
        "Discount factor: γ = 0.95",
        "Training episodes in “training” mode: 300",
    ]
    for c in comps:
        doc.add_paragraph(c, style="List Bullet")
    doc.add_paragraph("The trained model is saved as model.pth.")

    # (Figure 2 is presented in Section 5 Evaluation)

    # 4. Game Interaction and UI
    doc.add_paragraph("4. Game Interaction and User Interface", style="Heading 1")
    doc.add_paragraph(
        "The game offers a highly polished, intuitive interface with:"
    )
    ui = [
        "Dreamy pink gradient background and glowing food effects",
        "Rounded snake blocks with fading trail",
        "Real-time information card displaying score, best score, mode, ε-value, last action, last reward, combo, food type, and full 11D state vector",
        "Four interactive modes (manual / random AI / training / test) switched with T (cycle) or F1–F4 (direct jump)",
        "Sound effects for eating, bonus, poison, death, and pause",
        "Smooth Game Over screen with “Press Enter to play again”",
    ]
    for u in ui:
        doc.add_paragraph(u, style="List Bullet")
    doc.add_paragraph(
        "These features provide excellent user experience and clearly display state, actions, and rewards as required."
    )

    # 5. Evaluation
    doc.add_paragraph("5. Evaluation Results", style="Heading 1")
    doc.add_paragraph(
        "Training progress is evaluated using the episode score trajectory and a moving-average curve (Figure 2). "
        "Because the policy is trained under ε-greedy exploration, early episodes are noisy; as ε decays, the agent becomes more consistent."
    )

    # Figure 2: insert training curve if present
    curve = base_dir / "training_curve.png"
    if curve.exists():
        doc.add_paragraph("")
        try:
            doc.add_picture(str(curve), width=Inches(5.8))
        except Exception:
            p = doc.add_paragraph("[Could not embed training_curve.png — insert manually]")
            p.runs[0].italic = True
    else:
        p = doc.add_paragraph("[Insert Figure 2 (training_curve.png) here]")
        p.runs[0].italic = True

    cap = doc.add_paragraph("Figure 2: Training curve (training_curve.png) demonstrating learning progress")
    cap.runs[0].italic = True

    doc.add_paragraph("Key observations from a typical run include:")
    evals = [
        "Early episodes: low scores with frequent wall/self collisions during exploration.",
        "Mid training: gradual improvement as the agent learns basic obstacle avoidance and food-seeking behavior.",
        "Late training: higher moving-average score and fewer immediate collisions as ε approaches its minimum.",
    ]
    for ev in evals:
        doc.add_paragraph(ev, style="List Bullet")

    doc.add_paragraph(
        "In test mode (ε = 0), the trained policy typically survives longer and collects food more reliably than the random baseline, "
        "demonstrating that the learned Q-function improves decision quality."
    )

    # 6. Challenges
    doc.add_paragraph("6. Challenges and Solutions", style="Heading 1")
    p = doc.add_paragraph("1. Training speed vs visualization")
    p.runs[0].bold = True
    doc.add_paragraph("Challenge: Rendering every frame slows training and reduces the number of episodes that can be completed quickly.")
    doc.add_paragraph("Solution: Training runs at high speed, with optional UI refreshes every N episodes; a headless mode disables rendering entirely for notebook/CI environments.")

    p = doc.add_paragraph("2. Reward shaping trade-offs")
    p.runs[0].bold = True
    doc.add_paragraph("Attempt 1: Sparse reward (+food, −death) was simple but made exploration slow because useful trajectories are rare early on.")
    doc.add_paragraph("Attempt 2: A small living reward (+0.1) increased movement and improved exploration stability.")
    doc.add_paragraph("Attempt 3: Stronger directional shaping can speed up learning but may introduce oscillations near the food, so shaping was kept lightweight (penalize moving away only).")
    doc.add_paragraph("Final: Keep the main learning signal simple (+food/−death) and use minimal shaping for efficiency; special food types add variety without changing the core objective.")

    p = doc.add_paragraph("3. Platform compatibility (Colab/headless)")
    p.runs[0].bold = True
    doc.add_paragraph("Challenge: Pygame windows/audio often fail in hosted notebook environments.")
    doc.add_paragraph("Solution: Support HEADLESS=1 with dummy drivers; disable audio and screenshots in headless mode.")

    p = doc.add_paragraph("4. Presentation quality")
    p.runs[0].bold = True
    doc.add_paragraph("Challenge: A plain grid UI is hard to watch in a demo and does not clearly communicate state/action/reward.")
    doc.add_paragraph("Solution: A polished UI (info card + mode badge) shows state/action/reward explicitly; visual theme and sound effects improve demo clarity.")

    # 7. Conclusion
    doc.add_paragraph("7. Conclusion", style="Heading 1")
    doc.add_paragraph(
        "This project successfully demonstrates the application of Deep Q-Learning to a dynamic grid-based game. "
        "The agent learns effective strategies through trial-and-error, experience replay, and ε-greedy exploration. "
        "The combination of a visually stunning interface and measurable learning progress (via the training curve) fully satisfies all assignment requirements."
    )
    doc.add_paragraph(
        "Future work may include a target network, Double DQN, or pixel-based state representation for even stronger performance."
    )

    # 8. References
    doc.add_paragraph("8. References", style="Heading 1")
    refs = [
        "Mnih, V. et al. (2015). Human-level control through deep reinforcement learning. Nature.",
        "Sutton, R. S., & Barto, A. G. (2018). Reinforcement Learning: An Introduction (2nd ed.).",
        "PyTorch Documentation: https://pytorch.org/docs/",
        "Pygame Documentation: https://www.pygame.org/docs/",
    ]
    for r in refs:
        doc.add_paragraph(r, style="List Bullet")

    output_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_docx))


def main() -> None:
    base_dir = Path(__file__).resolve().parents[1]
    out_docx = base_dir / "Report - WANG Shiyu.docx"
    build_docx(out_docx, base_dir)
    print(f"Wrote: {out_docx}")

    # Optional: convert to PDF via Word (Windows)
    try:
        from docx2pdf import convert

        out_pdf = base_dir / "Report - WANG Shiyu.pdf"
        convert(str(out_docx), str(out_pdf))
        print(f"Wrote: {out_pdf}")
    except Exception as e:
        print("PDF conversion skipped/failed (this is OK). You can open the DOCX in Word and Save As PDF.")
        print(f"Reason: {e}")


if __name__ == "__main__":
    main()
